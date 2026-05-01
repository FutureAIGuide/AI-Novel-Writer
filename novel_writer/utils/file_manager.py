"""File management utilities for saving and loading stories."""

from __future__ import annotations

import json
import os
from pathlib import Path

from novel_writer.models.story import Story


class FileManager:
    """Handles saving and loading Story objects as JSON files."""

    def __init__(self, output_dir: str = "stories") -> None:
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _story_path(self, title: str) -> Path:
        """Return the file path for a story with the given title."""
        safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in title)
        safe_title = safe_title.strip().replace(" ", "_")
        return self.output_dir / f"{safe_title}.json"

    def save(self, story: Story) -> Path:
        """Serialize a Story to a JSON file and return the file path.

        Args:
            story: The Story instance to save.

        Returns:
            The Path of the saved file.
        """
        path = self._story_path(story.title)
        with open(path, "w", encoding="utf-8") as fh:
            fh.write(story.model_dump_json(indent=2))
        return path

    def load(self, title: str) -> Story:
        """Load a Story from a JSON file by title.

        Args:
            title: The story title (used to locate the file).

        Returns:
            A Story instance loaded from disk.

        Raises:
            FileNotFoundError: If no saved story with that title exists.
        """
        path = self._story_path(title)
        if not path.exists():
            raise FileNotFoundError(f"No saved story found for title: '{title}' at {path}")
        with open(path, "r", encoding="utf-8") as fh:
            data = json.load(fh)
        return Story(**data)

    def load_from_path(self, path: str | Path) -> Story:
        """Load a Story from an explicit file path.

        Args:
            path: Path to the JSON file.

        Returns:
            A Story instance.
        """
        with open(path, "r", encoding="utf-8") as fh:
            data = json.load(fh)
        return Story(**data)

    def list_stories(self) -> list[str]:
        """Return a list of saved story titles (without extension).

        Returns:
            List of story titles found in the output directory.
        """
        return [p.stem.replace("_", " ") for p in sorted(self.output_dir.glob("*.json"))]

    def export_text(self, story: Story) -> Path:
        """Export the full novel as a plain-text file.

        Args:
            story: The Story instance.

        Returns:
            The Path of the exported text file.
        """
        safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in story.title)
        safe_title = safe_title.strip().replace(" ", "_")
        path = self.output_dir / f"{safe_title}.txt"
        with open(path, "w", encoding="utf-8") as fh:
            fh.write(f"{story.title.upper()}\n")
            fh.write("=" * len(story.title) + "\n\n")
            if story.premise:
                fh.write(f"Premise: {story.premise}\n\n")
            for chapter in story.chapters:
                fh.write(f"\n{chapter.display_title()}\n")
                fh.write("-" * len(chapter.display_title()) + "\n\n")
                fh.write(chapter.content)
                fh.write("\n\n")
        return path

    def export_markdown(self, story: Story) -> Path:
        """Export the full novel as Markdown."""
        safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in story.title)
        safe_title = safe_title.strip().replace(" ", "_")
        path = self.output_dir / f"{safe_title}.md"
        lines: list[str] = [
            f"# {story.title}",
            "",
        ]
        if story.genre:
            lines.append(f"**Genre:** {story.genre}")
            lines.append("")
        if story.premise:
            lines.append("## Premise")
            lines.append("")
            lines.append(story.premise)
            lines.append("")
        if story.setting:
            lines.append("## Setting")
            lines.append("")
            lines.append(story.setting)
            lines.append("")
        if story.outline:
            lines.append("## Outline")
            lines.append("")
            lines.append(story.outline)
            lines.append("")
        for chapter in story.chapters:
            title = chapter.title or f"Chapter {chapter.number}"
            lines.append(f"## {chapter.display_title()}")
            lines.append("")
            if chapter.summary:
                lines.append(f"*{chapter.summary}*")
                lines.append("")
            lines.append(chapter.content or "_No content._")
            lines.append("")
        path.write_text("\n".join(lines), encoding="utf-8")
        return path

    def export_docx(self, story: Story) -> Path:
        """Export the manuscript as ``.docx`` (requires ``python-docx``)."""
        try:
            from docx import Document  # type: ignore[import]
        except ImportError as exc:
            raise ImportError(
                "Export to DOCX requires python-docx. Install with: pip install python-docx"
            ) from exc

        safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in story.title)
        safe_title = safe_title.strip().replace(" ", "_")
        path = self.output_dir / f"{safe_title}.docx"

        document = Document()
        document.add_heading(story.title, level=0)
        if story.genre:
            document.add_paragraph(f"Genre: {story.genre}")
        if story.premise:
            document.add_heading("Premise", level=1)
            document.add_paragraph(story.premise)
        for chapter in story.chapters:
            document.add_heading(chapter.display_title(), level=1)
            if chapter.summary:
                document.add_paragraph(chapter.summary, style="Intense Quote")
            document.add_paragraph(chapter.content or "")

        document.save(str(path))
        return path
